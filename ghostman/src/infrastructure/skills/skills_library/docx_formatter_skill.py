"""
DOCX Formatter Skill - Reformat and standardize DOCX documents.

This skill provides automated formatting operations for Microsoft Word (.docx)
documents, including font standardization, margin correction, spacing normalization,
bullet list fixes, spell checking, case correction, and heading normalization.

The original file is never overwritten; a new file with the suffix "_formatted"
is created in the same directory.
"""

import logging
import re
from typing import List, Any, Dict, Optional
from pathlib import Path

from ..interfaces.base_skill import (
    BaseSkill,
    SkillMetadata,
    SkillParameter,
    SkillResult,
    PermissionType,
    SkillCategory,
)

logger = logging.getLogger("ghostman.skills.docx_formatter")

# All available formatting operations in execution order
ALL_OPERATIONS = [
    "standardize_fonts",
    "fix_margins",
    "normalize_spacing",
    "fix_bullets",
    "fix_spelling",
    "fix_case",
    "normalize_headings",
]


class DocxFormatterSkill(BaseSkill):
    """
    Skill for reformatting and standardizing DOCX documents.

    Applies a configurable set of formatting operations to a Word document,
    producing a new formatted copy. Operations include font standardization,
    margin correction, spacing normalization, bullet list fixes, spell
    checking, case correction, and heading normalization.

    Requirements:
        - python-docx for document manipulation
        - pyspellchecker for spell checking (optional, only for fix_spelling)

    Example:
        >>> skill = DocxFormatterSkill()
        >>> result = await skill.execute(
        ...     file_path="C:/Documents/report.docx",
        ...     operations=["standardize_fonts", "fix_margins"]
        ... )
        >>> print(result.data["formatted_path"])
        "C:\\Documents\\report_formatted.docx"
    """

    @property
    def metadata(self) -> SkillMetadata:
        """Return skill metadata."""
        return SkillMetadata(
            skill_id="docx_formatter",
            name="Document Formatter",
            description="Reformat and standardize DOCX documents with font, margin, spacing, and spelling fixes",
            category=SkillCategory.FILE_MANAGEMENT,
            icon="\U0001f4c4",
            enabled_by_default=True,
            requires_confirmation=True,
            permissions_required=[PermissionType.FILE_READ, PermissionType.FILE_WRITE],
            ai_callable=True,
            version="1.0.0",
            author="Ghostman",
        )

    @property
    def parameters(self) -> List[SkillParameter]:
        """Return list of parameters this skill accepts."""
        return [
            SkillParameter(
                name="file_path",
                type=str,
                required=True,
                description="Path to the DOCX file to format",
                constraints={"min_length": 1, "max_length": 500},
            ),
            SkillParameter(
                name="operations",
                type=list,
                required=False,
                description=(
                    "List of formatting operations to apply. Available: "
                    "standardize_fonts, fix_margins, normalize_spacing, "
                    "fix_bullets, fix_spelling, fix_case, normalize_headings"
                ),
                constraints={
                    "items_type": str,
                    "choices": ALL_OPERATIONS,
                },
            ),
        ]

    async def execute(self, **params: Any) -> SkillResult:
        """
        Execute the document formatting skill.

        Validates the input file, loads the document, runs the requested
        formatting operations in order, and saves a new formatted copy.

        Args:
            **params: Validated parameters (file_path, operations)

        Returns:
            SkillResult with formatting summary, file paths, and change counts
        """
        try:
            # --- Dependency check ---
            try:
                from docx import Document
                from docx.shared import Pt, Inches  # noqa: F401
            except ImportError:
                return SkillResult(
                    success=False,
                    message="python-docx is not installed",
                    error=(
                        "The python-docx package is required for document formatting. "
                        "Install it with: pip install python-docx"
                    ),
                )

            # --- File validation ---
            file_path = Path(params["file_path"])

            if not file_path.exists():
                return SkillResult(
                    success=False,
                    message=f"File not found: {file_path}",
                    error=f"The file does not exist: {file_path}",
                )

            if file_path.suffix.lower() != ".docx":
                return SkillResult(
                    success=False,
                    message=f"Not a DOCX file: {file_path.name}",
                    error="Only .docx files are supported. Please provide a valid Word document.",
                )

            # --- Load document ---
            logger.info(f"Loading document: {file_path}")
            doc = Document(str(file_path))

            # --- Determine operations ---
            operations = params.get("operations")
            if operations is None:
                # Lazy import to avoid circular imports at module level
                from ...storage.settings_manager import settings

                operations = settings.get(
                    "tools.docx_formatter.default_operations",
                    list(ALL_OPERATIONS),
                )

            # Validate requested operations
            invalid_ops = [op for op in operations if op not in ALL_OPERATIONS]
            if invalid_ops:
                return SkillResult(
                    success=False,
                    message=f"Invalid operations: {', '.join(invalid_ops)}",
                    error=f"Valid operations are: {', '.join(ALL_OPERATIONS)}",
                )

            # --- Read formatting defaults from settings ---
            from ...storage.settings_manager import settings

            font_name = settings.get("tools.docx_formatter.default_font", "Calibri")
            font_size = settings.get("tools.docx_formatter.default_font_size", 11)
            line_spacing = settings.get("tools.docx_formatter.line_spacing", 1.15)
            margins = settings.get(
                "tools.docx_formatter.margins",
                {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
            )

            # --- Execute operations in order ---
            changes: Dict[str, Any] = {}

            for operation in ALL_OPERATIONS:
                if operation not in operations:
                    continue

                try:
                    if operation == "standardize_fonts":
                        changes[operation] = self._standardize_fonts(doc, font_name, font_size)
                    elif operation == "fix_margins":
                        changes[operation] = self._fix_margins(doc, margins)
                    elif operation == "normalize_spacing":
                        changes[operation] = self._normalize_spacing(doc, line_spacing)
                    elif operation == "fix_bullets":
                        changes[operation] = self._fix_bullets(doc)
                    elif operation == "fix_spelling":
                        changes[operation] = self._fix_spelling(doc)
                    elif operation == "fix_case":
                        changes[operation] = self._fix_case(doc)
                    elif operation == "normalize_headings":
                        changes[operation] = self._normalize_headings(doc)

                    logger.info(f"Operation '{operation}' completed: {changes[operation]}")

                except Exception as e:
                    logger.error(f"Operation '{operation}' failed: {e}", exc_info=True)
                    changes[operation] = f"error: {e}"

            # --- Save formatted document (never overwrite original) ---
            formatted_path = file_path.parent / f"{file_path.stem}_formatted.docx"
            doc.save(str(formatted_path))
            logger.info(f"Formatted document saved: {formatted_path}")

            # --- Compute total numeric changes ---
            total_changes = 0
            for value in changes.values():
                if isinstance(value, int):
                    total_changes += value
                elif isinstance(value, bool) and value:
                    total_changes += 1

            result_data = {
                "original_path": str(file_path),
                "formatted_path": str(formatted_path),
                "changes": changes,
                "total_changes": total_changes,
            }

            operations_summary = ", ".join(operations)
            return SkillResult(
                success=True,
                message=(
                    f"Document formatted successfully with {total_changes} changes "
                    f"({operations_summary}). Saved to: {formatted_path.name}"
                ),
                data=result_data,
                action_taken=f"Formatted {file_path.name} with operations: {operations_summary}",
            )

        except Exception as e:
            logger.error(f"Document formatting failed: {e}", exc_info=True)
            return SkillResult(
                success=False,
                message="Document formatting failed",
                error=str(e),
            )

    # ------------------------------------------------------------------
    # Private operation methods
    # ------------------------------------------------------------------

    def _standardize_fonts(self, doc: Any, font_name: str, font_size: int) -> int:
        """
        Standardize all fonts in the document to the specified font and size.

        Iterates every run in every paragraph and sets a uniform font name
        and point size.

        Args:
            doc: python-docx Document object
            font_name: Target font family name (e.g. "Calibri")
            font_size: Target font size in points (e.g. 11)

        Returns:
            Number of runs changed
        """
        from docx.shared import Pt

        changes_count = 0

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                changes_count += 1

        logger.debug(f"Standardized fonts: {changes_count} runs updated to {font_name} {font_size}pt")
        return changes_count

    def _fix_margins(self, doc: Any, margins: Dict[str, float]) -> bool:
        """
        Set all section margins to the specified values.

        Args:
            doc: python-docx Document object
            margins: Dict with 'top', 'bottom', 'left', 'right' values in inches

        Returns:
            True if margins were applied to at least one section
        """
        from docx.shared import Inches

        applied = False

        for section in doc.sections:
            section.top_margin = Inches(margins.get("top", 1.0))
            section.bottom_margin = Inches(margins.get("bottom", 1.0))
            section.left_margin = Inches(margins.get("left", 1.0))
            section.right_margin = Inches(margins.get("right", 1.0))
            applied = True

        logger.debug(f"Fixed margins for {len(doc.sections)} section(s)")
        return applied

    def _normalize_spacing(self, doc: Any, line_spacing: float) -> int:
        """
        Normalize line spacing and paragraph spacing throughout the document.

        Sets uniform line spacing on all paragraphs. For body text (non-heading
        paragraphs), space_before and space_after are set to zero.

        Args:
            doc: python-docx Document object
            line_spacing: Line spacing multiplier (e.g. 1.15)

        Returns:
            Number of paragraphs changed
        """
        from docx.shared import Pt

        changes_count = 0

        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing = line_spacing

            # Only zero out spacing for body text, not headings
            style_name = paragraph.style.name if paragraph.style else ""
            is_heading = style_name.lower().startswith("heading")

            if not is_heading:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

            changes_count += 1

        logger.debug(f"Normalized spacing: {changes_count} paragraphs updated")
        return changes_count

    def _fix_bullets(self, doc: Any) -> int:
        """
        Detect and normalize bullet/list paragraph indentation.

        Identifies bullet paragraphs by style name ('List' patterns) and by
        leading bullet characters. Normalizes first-level indentation to
        0.5 inches and nested bullets to 1.0 inches.

        Args:
            doc: python-docx Document object

        Returns:
            Number of bullet paragraphs fixed
        """
        from docx.shared import Inches

        changes_count = 0
        bullet_chars = {"\u2022", "-", "*", "\u25e6"}  # bullet, dash, asterisk, white bullet
        nested_chars = {"\u25e6"}  # white bullet indicates nested

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""
            text = paragraph.text.strip()

            is_list_style = "List" in style_name
            starts_with_bullet = any(text.startswith(ch) for ch in bullet_chars) if text else False

            if not is_list_style and not starts_with_bullet:
                continue

            # Determine nesting level
            is_nested = False

            # Check for double indentation (existing left_indent >= 1 inch)
            current_indent = paragraph.paragraph_format.left_indent
            if current_indent is not None:
                # left_indent is in EMU; 1 inch = 914400 EMU
                if current_indent >= 914400:
                    is_nested = True

            # Check for nested bullet characters
            if any(text.startswith(ch) for ch in nested_chars):
                is_nested = True

            # Check style for nesting (e.g. "List Bullet 2", "List Number 2")
            if re.search(r"List\s+\w+\s+[2-9]", style_name):
                is_nested = True

            if is_nested:
                paragraph.paragraph_format.left_indent = Inches(1.0)
            else:
                paragraph.paragraph_format.left_indent = Inches(0.5)

            changes_count += 1

        logger.debug(f"Fixed bullets: {changes_count} paragraphs normalized")
        return changes_count

    def _fix_spelling(self, doc: Any) -> int:
        """
        Run spell checking on all document text and apply confident corrections.

        Uses the pyspellchecker library to identify misspelled words and replace
        them with the top correction, operating at the run level to preserve
        formatting.

        Args:
            doc: python-docx Document object

        Returns:
            Number of words corrected
        """
        try:
            from spellchecker import SpellChecker
        except ImportError:
            logger.warning(
                "pyspellchecker not installed. Skipping spell check. "
                "Install with: pip install pyspellchecker"
            )
            return 0

        spell = SpellChecker()
        changes_count = 0

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if not run.text or not run.text.strip():
                    continue

                original_text = run.text
                words = re.findall(r"\b[a-zA-Z]+\b", original_text)

                if not words:
                    continue

                # Build lowercase→original mapping since spell.unknown() lowercases
                lower_to_original = {}
                for w in words:
                    lower_to_original.setdefault(w.lower(), w)

                misspelled = spell.unknown(words)

                if not misspelled:
                    continue

                new_text = original_text
                for lower_word in misspelled:
                    correction = spell.correction(lower_word)

                    if correction is None or correction == lower_word:
                        continue

                    # Look up original-cased word for pattern matching
                    orig_word = lower_to_original.get(lower_word, lower_word)

                    # Preserve original casing pattern when applying correction
                    if orig_word[0].isupper() and not orig_word.isupper():
                        correction = correction.capitalize()
                    elif orig_word.isupper():
                        correction = correction.upper()

                    # Replace the word in the run text using word boundaries
                    # Use re.IGNORECASE to match regardless of case variations
                    new_text = re.sub(
                        r"\b" + re.escape(orig_word) + r"\b",
                        correction,
                        new_text,
                    )
                    changes_count += 1

                if new_text != original_text:
                    run.text = new_text

        logger.debug(f"Fixed spelling: {changes_count} words corrected")
        return changes_count

    def _fix_case(self, doc: Any) -> int:
        """
        Detect and correct words with bad casing.

        Identifies two patterns:
        1. Mixed bad casing: lowercase first letter followed by uppercase
           (e.g. "wORD", "hELLO") - converted to lowercase.
        2. Unnecessary ALL CAPS: words longer than 3 characters that are
           entirely uppercase but are not likely acronyms - converted to
           title case.

        Operates at the run level to preserve formatting.

        Args:
            doc: python-docx Document object

        Returns:
            Number of words fixed
        """
        changes_count = 0

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if not run.text or not run.text.strip():
                    continue

                original_text = run.text
                words = re.findall(r"\b[a-zA-Z]+\b", original_text)

                if not words:
                    continue

                new_text = original_text

                for word in words:
                    if len(word) <= 1:
                        continue

                    fixed_word = None

                    # Pattern 1: lowercase first char, uppercase chars in the rest
                    # e.g. "wORD", "hELLO" — but not normal camelCase like "iPhone"
                    if re.search(r"^[a-z][A-Z]", word):
                        # Check that this is genuinely bad casing, not intentional
                        # camelCase. Bad casing has mostly uppercase after the first char.
                        upper_count = sum(1 for c in word[1:] if c.isupper())
                        if upper_count > len(word[1:]) / 2:
                            fixed_word = word.lower()

                    # Pattern 2: ALL CAPS word that is too long to be an acronym
                    elif len(word) > 3 and word.isupper():
                        fixed_word = word.title()

                    if fixed_word and fixed_word != word:
                        new_text = re.sub(
                            r"\b" + re.escape(word) + r"\b",
                            fixed_word,
                            new_text,
                            count=1,
                        )
                        changes_count += 1

                if new_text != original_text:
                    run.text = new_text

        logger.debug(f"Fixed case: {changes_count} words corrected")
        return changes_count

    def _normalize_headings(self, doc: Any) -> int:
        """
        Normalize heading paragraphs to consistent font sizes and bold style.

        Enforces:
            - Heading 1: 16pt, bold
            - Heading 2: 14pt, bold
            - Heading 3: 12pt, bold

        Args:
            doc: python-docx Document object

        Returns:
            Number of headings normalized
        """
        from docx.shared import Pt

        heading_sizes = {
            "Heading 1": 16,
            "Heading 2": 14,
            "Heading 3": 12,
        }

        changes_count = 0

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""

            if style_name not in heading_sizes:
                continue

            target_size = heading_sizes[style_name]

            for run in paragraph.runs:
                run.font.size = Pt(target_size)
                run.font.bold = True

            changes_count += 1

        logger.debug(f"Normalized headings: {changes_count} headings updated")
        return changes_count

    # ------------------------------------------------------------------
    # Lifecycle hooks
    # ------------------------------------------------------------------

    async def on_success(self, result: SkillResult) -> None:
        """Log successful document formatting."""
        logger.info(f"Document formatter succeeded: {result.message}")

    async def on_error(self, result: SkillResult) -> None:
        """Log document formatting failure."""
        logger.warning(f"Document formatter failed: {result.error}")
