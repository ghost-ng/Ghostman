"""
DOCX Formatter Skill - Reformat and standardize DOCX documents.

This skill provides automated formatting operations for Microsoft Word (.docx)
documents, including font standardization, margin correction, spacing normalization,
bullet list fixes, spell checking, case correction, and heading normalization.

Supports two modes:
- **One-shot**: Applies operations and saves to ``<name>_formatted.docx``.
- **Session**: Copies file to a temp location, applies operations in-place
  across multiple turns, then saves a final copy on session end.
"""

import ast
import logging
import os
import re
import shutil
import tempfile
import threading
from dataclasses import dataclass
from typing import List, Any, Dict, Optional
from pathlib import Path

from ..interfaces.base_skill import (
    BaseSkill,
    SkillMetadata,
    SkillParameter,
    SkillResult,
    PermissionType,
    SkillCategory,
)

logger = logging.getLogger("specter.skills.docx_formatter")

# All available formatting operations in execution order
ALL_OPERATIONS = [
    "standardize_fonts",
    "fix_margins",
    "normalize_spacing",
    "fix_bullets",
    "fix_spelling",
    "fix_case",
    "normalize_headings",
    "find_replace",
    "set_font_color",
    "set_alignment",
    "set_indent",
]

# Named color → hex mapping for set_font_color
COLOR_NAME_MAP = {
    "blue": "#0000FF",
    "red": "#FF0000",
    "green": "#008000",
    "black": "#000000",
    "white": "#FFFFFF",
    "gray": "#808080",
    "grey": "#808080",
    "orange": "#FFA500",
    "purple": "#800080",
    "navy": "#000080",
    "teal": "#008080",
    "maroon": "#800000",
    "yellow": "#FFFF00",
}


@dataclass
class CodeValidationResult:
    """Result from AST-based code validation."""
    valid: bool
    error: str = ""
    line_count: int = 0
    char_count: int = 0


@dataclass
class CodeExecutionResult:
    """Result from sandboxed exec() of AI-generated code."""
    success: bool
    message: str = ""
    error: str = ""
    paragraphs_before: int = 0
    runs_before: int = 0
    paragraphs_after: int = 0
    runs_after: int = 0


class DocxFormatterSkill(BaseSkill):
    """
    Skill for reformatting and standardizing DOCX documents.

    Applies a configurable set of formatting operations to a Word document,
    producing a new formatted copy. Operations include font standardization,
    margin correction, spacing normalization, bullet list fixes, spell
    checking, case correction, and heading normalization.

    Requirements:
        - python-docx for document manipulation
        - pyspellchecker for spell checking (optional, only for fix_spelling)

    Example:
        >>> skill = DocxFormatterSkill()
        >>> result = await skill.execute(
        ...     file_path="C:/Documents/report.docx",
        ...     operations=["standardize_fonts", "fix_margins"]
        ... )
        >>> print(result.data["formatted_path"])
        "C:\\Documents\\report_formatted.docx"
    """

    @property
    def metadata(self) -> SkillMetadata:
        """Return skill metadata."""
        return SkillMetadata(
            skill_id="docx_formatter",
            name="Document Formatter",
            description="Reformat and standardize DOCX documents with font, margin, spacing, and spelling fixes",
            category=SkillCategory.FILE_MANAGEMENT,
            icon="\U0001f4c4",
            enabled_by_default=True,
            requires_confirmation=True,
            permissions_required=[PermissionType.FILE_READ, PermissionType.FILE_WRITE],
            ai_callable=True,
            version="1.0.0",
            author="Specter",
        )

    @property
    def parameters(self) -> List[SkillParameter]:
        """Return list of parameters this skill accepts."""
        return [
            SkillParameter(
                name="file_path",
                type=str,
                required=True,
                description="Path to the DOCX file to format",
                constraints={"min_length": 1, "max_length": 500},
            ),
            SkillParameter(
                name="operations",
                type=list,
                required=False,
                description=(
                    "List of formatting operations to apply. Available: "
                    "standardize_fonts, fix_margins, normalize_spacing, "
                    "fix_bullets, fix_spelling, fix_case, normalize_headings, "
                    "find_replace, set_font_color, set_alignment, set_indent"
                ),
                constraints={
                    "items_type": str,
                    "choices": ALL_OPERATIONS,
                },
            ),
            SkillParameter(
                name="font_size",
                type=int,
                required=False,
                description="Font size in points (overrides default). Example: 14",
                constraints={"min": 6, "max": 72},
            ),
            SkillParameter(
                name="font_name",
                type=str,
                required=False,
                description="Font name to use (overrides default). Example: Arial",
            ),
            SkillParameter(
                name="find_text",
                type=str,
                required=False,
                description="Text to find (used with find_replace operation)",
            ),
            SkillParameter(
                name="replace_text",
                type=str,
                required=False,
                description="Replacement text (used with find_replace operation)",
            ),
            SkillParameter(
                name="font_color",
                type=str,
                required=False,
                description="Font color as hex '#0000FF' or name 'blue' (used with set_font_color)",
            ),
            SkillParameter(
                name="target_scope",
                type=str,
                required=False,
                description="Scope for color/alignment operations: 'all', 'headings', or 'body'",
                constraints={"choices": ["all", "headings", "body"]},
            ),
            SkillParameter(
                name="alignment",
                type=str,
                required=False,
                description="Paragraph alignment: 'left', 'center', 'right', or 'justify'",
                constraints={"choices": ["left", "center", "right", "justify"]},
            ),
            SkillParameter(
                name="indent_inches",
                type=float,
                required=False,
                description="First-line indent in inches (used with set_indent). Example: 0.5",
                constraints={"min": 0.0, "max": 3.0},
            ),
            SkillParameter(
                name="bullet_indent",
                type=float,
                required=False,
                description="Indent per bullet level in inches (default 0.5)",
                constraints={"min": 0.1, "max": 2.0},
            ),
        ]

    async def execute(self, **params: Any) -> SkillResult:
        """
        Execute the document formatting skill.

        Validates the input file, loads the document, runs the requested
        formatting operations in order, and saves a new formatted copy.

        Args:
            **params: Validated parameters (file_path, operations)

        Returns:
            SkillResult with formatting summary, file paths, and change counts
        """
        try:
            # --- Dependency check ---
            try:
                from docx import Document
                from docx.shared import Pt, Inches  # noqa: F401
            except ImportError:
                return SkillResult(
                    success=False,
                    message="python-docx is not installed",
                    error=(
                        "The python-docx package is required for document formatting. "
                        "Install it with: pip install python-docx"
                    ),
                )

            # --- File validation ---
            file_path = Path(params["file_path"])

            if not file_path.exists():
                return SkillResult(
                    success=False,
                    message=f"File not found: {file_path}",
                    error=f"The file does not exist: {file_path}",
                )

            if file_path.suffix.lower() != ".docx":
                return SkillResult(
                    success=False,
                    message=f"Not a DOCX file: {file_path.name}",
                    error="Only .docx files are supported. Please provide a valid Word document.",
                )

            # --- Load document ---
            logger.info(f"Loading document: {file_path}")
            doc = Document(str(file_path))

            # --- Determine operations ---
            operations = params.get("operations")
            if operations is None:
                # Lazy import to avoid circular imports at module level
                from ...storage.settings_manager import settings

                operations = settings.get(
                    "tools.docx_formatter.default_operations",
                    list(ALL_OPERATIONS),
                )

            # Validate requested operations
            invalid_ops = [op for op in operations if op not in ALL_OPERATIONS]
            if invalid_ops:
                return SkillResult(
                    success=False,
                    message=f"Invalid operations: {', '.join(invalid_ops)}",
                    error=f"Valid operations are: {', '.join(ALL_OPERATIONS)}",
                )

            # --- Read formatting defaults from settings, allow param overrides ---
            from ...storage.settings_manager import settings

            font_name = params.get("font_name") or settings.get("tools.docx_formatter.default_font", "Calibri")
            font_size = params.get("font_size") or settings.get("tools.docx_formatter.default_font_size", 11)
            line_spacing = settings.get("tools.docx_formatter.line_spacing", 1.15)
            margins = settings.get(
                "tools.docx_formatter.margins",
                {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
            )
            bullet_indent = params.get("bullet_indent", 0.5)

            # --- Execute operations in order ---
            changes: Dict[str, Any] = {}

            for operation in ALL_OPERATIONS:
                if operation not in operations:
                    continue

                try:
                    if operation == "standardize_fonts":
                        changes[operation] = self._standardize_fonts(doc, font_name, font_size)
                    elif operation == "fix_margins":
                        changes[operation] = self._fix_margins(doc, margins)
                    elif operation == "normalize_spacing":
                        changes[operation] = self._normalize_spacing(doc, line_spacing)
                    elif operation == "fix_bullets":
                        changes[operation] = self._fix_bullets(doc, bullet_indent)
                    elif operation == "fix_spelling":
                        changes[operation] = self._fix_spelling(doc)
                    elif operation == "fix_case":
                        changes[operation] = self._fix_case(doc)
                    elif operation == "normalize_headings":
                        changes[operation] = self._normalize_headings(doc)
                    elif operation == "find_replace":
                        find_text = params.get("find_text", "")
                        replace_text = params.get("replace_text", "")
                        if find_text:
                            changes[operation] = self._find_replace(doc, find_text, replace_text)
                        else:
                            changes[operation] = "skipped: no find_text provided"
                    elif operation == "set_font_color":
                        font_color = params.get("font_color", "")
                        target_scope = params.get("target_scope", "all")
                        if font_color:
                            changes[operation] = self._set_font_color(doc, font_color, target_scope)
                        else:
                            changes[operation] = "skipped: no font_color provided"
                    elif operation == "set_alignment":
                        alignment = params.get("alignment", "")
                        target_scope = params.get("target_scope", "all")
                        if alignment:
                            changes[operation] = self._set_alignment(doc, alignment, target_scope)
                        else:
                            changes[operation] = "skipped: no alignment provided"
                    elif operation == "set_indent":
                        indent_inches = params.get("indent_inches")
                        if indent_inches is not None:
                            changes[operation] = self._set_indent(doc, indent_inches)
                        else:
                            changes[operation] = "skipped: no indent_inches provided"

                    logger.info(f"Operation '{operation}' completed: {changes[operation]}")

                except Exception as e:
                    logger.error(f"Operation '{operation}' failed: {e}", exc_info=True)
                    changes[operation] = f"error: {e}"

            # --- Save formatted document ---
            session_mode = params.get("_session_mode", False)
            if session_mode:
                # Session mode: save in-place (working on temp copy)
                doc.save(str(file_path))
                formatted_path = file_path
                logger.info(f"Session mode: saved in-place to {file_path}")
            else:
                # Normal mode: save to AppData temp dir
                formatted_path = self._get_output_path(file_path)
                if self._is_file_locked(formatted_path):
                    formatted_path = self._find_available_path(formatted_path)
                try:
                    doc.save(str(formatted_path))
                except PermissionError:
                    formatted_path = self._find_available_path(formatted_path)
                    doc.save(str(formatted_path))
                # Auto-open the formatted file
                self.open_file(str(formatted_path))
                logger.info(f"Formatted document saved: {formatted_path}")

            # --- Compute total numeric changes ---
            total_changes = 0
            for value in changes.values():
                if isinstance(value, int):
                    total_changes += value
                elif isinstance(value, bool) and value:
                    total_changes += 1

            result_data = {
                "original_path": str(file_path),
                "formatted_path": str(formatted_path),
                "changes": changes,
                "total_changes": total_changes,
            }

            operations_summary = ", ".join(operations)
            return SkillResult(
                success=True,
                message=(
                    f"Document formatted successfully with {total_changes} changes "
                    f"({operations_summary}). Saved to: {formatted_path.name}"
                ),
                data=result_data,
                action_taken=f"Formatted {file_path.name} with operations: {operations_summary}",
            )

        except Exception as e:
            logger.error(f"Document formatting failed: {e}", exc_info=True)
            return SkillResult(
                success=False,
                message="Document formatting failed",
                error=str(e),
            )

    # ------------------------------------------------------------------
    # Private operation methods
    # ------------------------------------------------------------------

    def _standardize_fonts(self, doc: Any, font_name: str, font_size: int) -> int:
        """
        Standardize all fonts in the document to the specified font and size.

        Iterates every run in every paragraph and sets a uniform font name
        and point size.

        Args:
            doc: python-docx Document object
            font_name: Target font family name (e.g. "Calibri")
            font_size: Target font size in points (e.g. 11)

        Returns:
            Number of runs changed
        """
        from docx.shared import Pt

        changes_count = 0

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                changes_count += 1

        logger.debug(f"Standardized fonts: {changes_count} runs updated to {font_name} {font_size}pt")
        return changes_count

    def _fix_margins(self, doc: Any, margins: Dict[str, float]) -> bool:
        """
        Set all section margins to the specified values.

        Args:
            doc: python-docx Document object
            margins: Dict with 'top', 'bottom', 'left', 'right' values in inches

        Returns:
            True if margins were applied to at least one section
        """
        from docx.shared import Inches

        applied = False

        for section in doc.sections:
            section.top_margin = Inches(margins.get("top", 1.0))
            section.bottom_margin = Inches(margins.get("bottom", 1.0))
            section.left_margin = Inches(margins.get("left", 1.0))
            section.right_margin = Inches(margins.get("right", 1.0))
            applied = True

        logger.debug(f"Fixed margins for {len(doc.sections)} section(s)")
        return applied

    def _normalize_spacing(self, doc: Any, line_spacing: float) -> int:
        """
        Normalize line spacing and paragraph spacing throughout the document.

        Sets uniform line spacing on all paragraphs. For body text (non-heading
        paragraphs), space_before and space_after are set to zero.

        Args:
            doc: python-docx Document object
            line_spacing: Line spacing multiplier (e.g. 1.15)

        Returns:
            Number of paragraphs changed
        """
        from docx.shared import Pt

        changes_count = 0

        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing = line_spacing

            # Only zero out spacing for body text, not headings
            style_name = paragraph.style.name if paragraph.style else ""
            is_heading = style_name.lower().startswith("heading")

            if not is_heading:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

            changes_count += 1

        logger.debug(f"Normalized spacing: {changes_count} paragraphs updated")
        return changes_count

    def _fix_bullets(self, doc: Any, base_indent: float = 0.5) -> int:
        """
        Detect and normalize bullet/list paragraph indentation with multi-level support.

        Identifies bullet paragraphs by style name ('List' patterns) and by
        leading bullet characters. Detects nesting level and applies incremental
        indentation: level * base_indent inches.

        Args:
            doc: python-docx Document object
            base_indent: Indent per bullet level in inches (default 0.5)

        Returns:
            Number of bullet paragraphs fixed
        """
        from docx.shared import Inches

        changes_count = 0

        for paragraph in doc.paragraphs:
            level = self._detect_bullet_level(paragraph)
            if level == 0:
                continue

            paragraph.paragraph_format.left_indent = Inches(level * base_indent)
            changes_count += 1

        logger.debug(f"Fixed bullets: {changes_count} paragraphs normalized (base_indent={base_indent}\")")
        return changes_count

    def _detect_bullet_level(self, paragraph: Any) -> int:
        """
        Detect the bullet/list nesting level of a paragraph.

        Returns:
            0 if not a bullet paragraph, 1+ for nesting level.
        """
        style_name = paragraph.style.name if paragraph.style else ""
        text = paragraph.text.strip()
        bullet_chars = {"\u2022", "-", "*", "\u25e6", "\u25aa", "\u25ab"}

        is_list_style = "List" in style_name
        starts_with_bullet = any(text.startswith(ch) for ch in bullet_chars) if text else False

        if not is_list_style and not starts_with_bullet:
            return 0

        # Priority 1: Style name number — "List Bullet 2" → level 2
        level_match = re.search(r"(\d+)", style_name) if is_list_style else None
        if level_match:
            return int(level_match.group(1))

        # Priority 2: Existing indent depth → estimate level
        current_indent = paragraph.paragraph_format.left_indent
        if current_indent is not None and current_indent > 0:
            # EMU per inch = 914400
            estimated_level = max(1, round(current_indent / (914400 * 0.5)))
            return estimated_level

        # Default: level 1
        return 1

    def _fix_spelling(self, doc: Any) -> int:
        """
        Run spell checking on all document text and apply confident corrections.

        Uses the pyspellchecker library to identify misspelled words and replace
        them with the top correction, operating at the run level to preserve
        formatting.

        Args:
            doc: python-docx Document object

        Returns:
            Number of words corrected
        """
        try:
            from spellchecker import SpellChecker
        except ImportError:
            logger.warning(
                "pyspellchecker not installed. Skipping spell check. "
                "Install with: pip install pyspellchecker"
            )
            return 0

        spell = SpellChecker()
        changes_count = 0

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if not run.text or not run.text.strip():
                    continue

                original_text = run.text
                words = re.findall(r"\b[a-zA-Z]+\b", original_text)

                if not words:
                    continue

                # Build lowercase→original mapping since spell.unknown() lowercases
                lower_to_original = {}
                for w in words:
                    lower_to_original.setdefault(w.lower(), w)

                misspelled = spell.unknown(words)

                if not misspelled:
                    continue

                new_text = original_text
                for lower_word in misspelled:
                    correction = spell.correction(lower_word)

                    if correction is None or correction == lower_word:
                        continue

                    # Look up original-cased word for pattern matching
                    orig_word = lower_to_original.get(lower_word, lower_word)

                    # Preserve original casing pattern when applying correction
                    if orig_word[0].isupper() and not orig_word.isupper():
                        correction = correction.capitalize()
                    elif orig_word.isupper():
                        correction = correction.upper()

                    # Replace the word in the run text using word boundaries
                    # Use re.IGNORECASE to match regardless of case variations
                    new_text = re.sub(
                        r"\b" + re.escape(orig_word) + r"\b",
                        correction,
                        new_text,
                    )
                    changes_count += 1

                if new_text != original_text:
                    run.text = new_text

        logger.debug(f"Fixed spelling: {changes_count} words corrected")
        return changes_count

    def _fix_case(self, doc: Any) -> int:
        """
        Detect and correct words with bad casing.

        Identifies two patterns:
        1. Mixed bad casing: lowercase first letter followed by uppercase
           (e.g. "wORD", "hELLO") - converted to lowercase.
        2. Unnecessary ALL CAPS: words longer than 3 characters that are
           entirely uppercase but are not likely acronyms - converted to
           title case.

        Operates at the run level to preserve formatting.

        Args:
            doc: python-docx Document object

        Returns:
            Number of words fixed
        """
        changes_count = 0

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if not run.text or not run.text.strip():
                    continue

                original_text = run.text
                words = re.findall(r"\b[a-zA-Z]+\b", original_text)

                if not words:
                    continue

                new_text = original_text

                for word in words:
                    if len(word) <= 1:
                        continue

                    fixed_word = None

                    # Pattern 1: lowercase first char, uppercase chars in the rest
                    # e.g. "wORD", "hELLO" — but not normal camelCase like "iPhone"
                    if re.search(r"^[a-z][A-Z]", word):
                        # Check that this is genuinely bad casing, not intentional
                        # camelCase. Bad casing has mostly uppercase after the first char.
                        upper_count = sum(1 for c in word[1:] if c.isupper())
                        if upper_count > len(word[1:]) / 2:
                            fixed_word = word.lower()

                    # Pattern 2: ALL CAPS word that is too long to be an acronym
                    elif len(word) > 3 and word.isupper():
                        fixed_word = word.title()

                    if fixed_word and fixed_word != word:
                        new_text = re.sub(
                            r"\b" + re.escape(word) + r"\b",
                            fixed_word,
                            new_text,
                            count=1,
                        )
                        changes_count += 1

                if new_text != original_text:
                    run.text = new_text

        logger.debug(f"Fixed case: {changes_count} words corrected")
        return changes_count

    def _normalize_headings(self, doc: Any) -> int:
        """
        Normalize heading paragraphs to consistent font sizes and bold style.

        Enforces:
            - Heading 1: 16pt, bold
            - Heading 2: 14pt, bold
            - Heading 3: 12pt, bold

        Args:
            doc: python-docx Document object

        Returns:
            Number of headings normalized
        """
        from docx.shared import Pt

        heading_sizes = {
            "Heading 1": 16,
            "Heading 2": 14,
            "Heading 3": 12,
        }

        changes_count = 0

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""

            if style_name not in heading_sizes:
                continue

            target_size = heading_sizes[style_name]

            for run in paragraph.runs:
                run.font.size = Pt(target_size)
                run.font.bold = True

            changes_count += 1

        logger.debug(f"Normalized headings: {changes_count} headings updated")
        return changes_count

    # ------------------------------------------------------------------
    # New operations: find_replace, set_font_color, set_alignment, set_indent
    # ------------------------------------------------------------------

    def _find_replace(self, doc: Any, find_text: str, replace_text: str) -> int:
        """
        Find and replace text throughout the document, preserving run-level formatting.

        First attempts run-level replacement to preserve formatting. Falls back to
        paragraph-level replacement for matches that span multiple runs.

        Args:
            doc: python-docx Document object
            find_text: Text to search for
            replace_text: Text to substitute

        Returns:
            Number of replacements made
        """
        changes_count = 0

        for paragraph in doc.paragraphs:
            # Try run-level replacement first (preserves formatting)
            for run in paragraph.runs:
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_text)
                    changes_count += run.text.count(replace_text)  # approximate

            # Fallback: check if the full paragraph text contains the target
            # (handles cases where find_text spans multiple runs)
            full_text = paragraph.text
            if find_text in full_text and changes_count == 0:
                # Rebuild the paragraph: keep first run's formatting, clear others
                if paragraph.runs:
                    new_text = full_text.replace(find_text, replace_text)
                    # Clear all runs then set text on the first run
                    first_run = paragraph.runs[0]
                    for run in paragraph.runs[1:]:
                        run.text = ""
                    first_run.text = new_text
                    changes_count += new_text.count(replace_text)

        logger.debug(f"Find/replace: '{find_text}' → '{replace_text}', {changes_count} replacement(s)")
        return changes_count

    def _resolve_color(self, color_input: str) -> Optional[tuple]:
        """
        Resolve a color name or hex string to (r, g, b) tuple.

        Args:
            color_input: Hex "#0000FF" or name "blue"

        Returns:
            (r, g, b) tuple or None if unrecognized
        """
        color_input = color_input.strip().lower()

        # Check named colors
        hex_val = COLOR_NAME_MAP.get(color_input)
        if not hex_val:
            # Check if it's a hex string
            if re.match(r"^#?[0-9a-f]{6}$", color_input):
                hex_val = color_input if color_input.startswith("#") else f"#{color_input}"
            else:
                return None

        hex_val = hex_val.lstrip("#")
        return (int(hex_val[0:2], 16), int(hex_val[2:4], 16), int(hex_val[4:6], 16))

    def _matches_scope(self, paragraph: Any, scope: str) -> bool:
        """Check if a paragraph matches the target scope."""
        style_name = paragraph.style.name if paragraph.style else ""
        is_heading = style_name.lower().startswith("heading")

        if scope == "all":
            return True
        elif scope == "headings":
            return is_heading
        elif scope == "body":
            return not is_heading
        return True

    def _set_font_color(self, doc: Any, font_color: str, target_scope: str = "all") -> int:
        """
        Set font color on runs matching the target scope.

        Args:
            doc: python-docx Document object
            font_color: Color as hex "#0000FF" or name "blue"
            target_scope: "all", "headings", or "body"

        Returns:
            Number of runs changed
        """
        from docx.shared import RGBColor

        rgb = self._resolve_color(font_color)
        if rgb is None:
            logger.warning(f"Unrecognized color: {font_color}")
            return 0

        changes_count = 0

        for paragraph in doc.paragraphs:
            if not self._matches_scope(paragraph, target_scope):
                continue

            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(*rgb)
                changes_count += 1

        logger.debug(f"Set font color: {font_color} on {target_scope}, {changes_count} run(s)")
        return changes_count

    def _set_alignment(self, doc: Any, alignment: str, target_scope: str = "all") -> int:
        """
        Set paragraph alignment on paragraphs matching the target scope.

        Args:
            doc: python-docx Document object
            alignment: "left", "center", "right", or "justify"
            target_scope: "all", "headings", or "body"

        Returns:
            Number of paragraphs changed
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }

        wd_align = align_map.get(alignment.lower())
        if wd_align is None:
            logger.warning(f"Unknown alignment: {alignment}")
            return 0

        changes_count = 0

        for paragraph in doc.paragraphs:
            if not self._matches_scope(paragraph, target_scope):
                continue

            paragraph.alignment = wd_align
            changes_count += 1

        logger.debug(f"Set alignment: {alignment} on {target_scope}, {changes_count} paragraph(s)")
        return changes_count

    def _set_indent(self, doc: Any, indent_inches: float) -> int:
        """
        Set first-line indent on body paragraphs (non-heading, non-list).

        Args:
            doc: python-docx Document object
            indent_inches: First-line indent in inches

        Returns:
            Number of paragraphs changed
        """
        from docx.shared import Inches

        changes_count = 0

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""
            is_heading = style_name.lower().startswith("heading")
            is_list = "List" in style_name

            # Only indent body paragraphs
            if is_heading or is_list:
                continue

            # Skip empty paragraphs
            if not paragraph.text.strip():
                continue

            paragraph.paragraph_format.first_line_indent = Inches(indent_inches)
            changes_count += 1

        logger.debug(f"Set indent: {indent_inches}\" on {changes_count} body paragraph(s)")
        return changes_count

    # ------------------------------------------------------------------
    # File-lock helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _get_appdata_temp_dir() -> Path:
        """Return %APPDATA%/Specter/temp/, creating it if needed."""
        appdata = os.environ.get("APPDATA", "")
        if appdata:
            d = Path(appdata) / "Specter" / "temp"
        else:
            d = Path(tempfile.gettempdir()) / "specter_sessions"
        d.mkdir(parents=True, exist_ok=True)
        return d

    @staticmethod
    def _get_output_path(original_path: Path) -> Path:
        """Return the output path inside AppData temp dir."""
        temp_dir = DocxFormatterSkill._get_appdata_temp_dir()
        return temp_dir / f"{original_path.stem}_formatted.docx"

    @staticmethod
    def _is_file_locked(path: Path) -> bool:
        """Check whether *path* is locked by another process (e.g. Word)."""
        if not path.exists():
            return False
        try:
            with open(path, "r+b"):
                return False
        except (PermissionError, OSError):
            return True

    @staticmethod
    def _find_available_path(base_path: Path) -> Path:
        """Try _formatted_2, _formatted_3, … until an unlocked name is found."""
        stem = base_path.stem.replace("_formatted", "")
        parent = base_path.parent
        for i in range(2, 20):
            candidate = parent / f"{stem}_formatted_{i}.docx"
            if not DocxFormatterSkill._is_file_locked(candidate):
                return candidate
        # Last resort — use a timestamped name
        import time
        ts = int(time.time())
        return parent / f"{stem}_formatted_{ts}.docx"

    @staticmethod
    def open_file(path: str) -> None:
        """Open a file with the system's default application."""
        try:
            p = Path(path)
            if not p.exists():
                logger.warning(f"Cannot open file — not found: {p}")
                return
            if os.name == "nt":
                os.startfile(str(p))
            else:
                import subprocess
                opener = "open" if os.uname().sysname == "Darwin" else "xdg-open"
                subprocess.Popen([opener, str(p)])
            logger.info(f"Opened file: {p}")
        except Exception as e:
            logger.warning(f"Failed to open file: {e}")

    # ------------------------------------------------------------------
    # AI code-generation sandbox
    # ------------------------------------------------------------------

    # Names that must never appear as ast.Name references
    _BLOCKED_NAMES = frozenset({
        "exec", "eval", "compile", "__import__", "open", "input",
        "breakpoint", "exit", "quit", "globals", "locals",
        "getattr", "setattr", "delattr", "vars", "dir",
        # Dangerous modules — block even as bare names
        "os", "sys", "subprocess", "pathlib", "shutil", "socket",
        "requests", "urllib", "pickle", "marshal", "ctypes",
        "importlib", "builtins", "__builtins__",
    })

    @staticmethod
    def _validate_code(code: str) -> CodeValidationResult:
        """
        Validate AI-generated python-docx code via AST analysis.

        Returns CodeValidationResult with valid=False on the first
        rule violation found.
        """
        line_count = code.count("\n") + 1
        char_count = len(code)

        if char_count > 5000:
            return CodeValidationResult(
                valid=False,
                error=f"Code too long: {char_count} chars (max 5000)",
                line_count=line_count, char_count=char_count,
            )
        if line_count > 100:
            return CodeValidationResult(
                valid=False,
                error=f"Code too long: {line_count} lines (max 100)",
                line_count=line_count, char_count=char_count,
            )

        try:
            tree = ast.parse(code)
        except SyntaxError as e:
            return CodeValidationResult(
                valid=False,
                error=f"Syntax error: {e}",
                line_count=line_count, char_count=char_count,
            )

        for node in ast.walk(tree):
            # --- Block forbidden imports ---
            if isinstance(node, ast.Import):
                for alias in node.names:
                    if alias.name.split(".")[0] != "docx":
                        return CodeValidationResult(
                            valid=False,
                            error=f"Forbidden import: 'import {alias.name}'",
                            line_count=line_count, char_count=char_count,
                        )

            if isinstance(node, ast.ImportFrom):
                module = node.module or ""
                if module.split(".")[0] != "docx":
                    return CodeValidationResult(
                        valid=False,
                        error=f"Forbidden import: 'from {module} import ...'",
                        line_count=line_count, char_count=char_count,
                    )
                for alias in node.names:
                    if alias.name == "*":
                        return CodeValidationResult(
                            valid=False,
                            error=f"Star import not allowed: 'from {module} import *'",
                            line_count=line_count, char_count=char_count,
                        )

            # --- Block dangerous name references ---
            if isinstance(node, ast.Name):
                if node.id in DocxFormatterSkill._BLOCKED_NAMES:
                    return CodeValidationResult(
                        valid=False,
                        error=f"Forbidden name: '{node.id}'",
                        line_count=line_count, char_count=char_count,
                    )

            # --- Block dunder attribute access (except __init__) ---
            if isinstance(node, ast.Attribute):
                attr = node.attr
                if attr.startswith("__") and attr.endswith("__") and attr != "__init__":
                    return CodeValidationResult(
                        valid=False,
                        error=f"Forbidden dunder access: '{attr}'",
                        line_count=line_count, char_count=char_count,
                    )

            # --- Block del statements ---
            if isinstance(node, ast.Delete):
                return CodeValidationResult(
                    valid=False,
                    error="'del' statements are not allowed",
                    line_count=line_count, char_count=char_count,
                )

            # --- Block .clear() calls ---
            if isinstance(node, ast.Call):
                func = node.func
                if isinstance(func, ast.Attribute) and func.attr == "clear":
                    return CodeValidationResult(
                        valid=False,
                        error="'.clear()' calls are not allowed",
                        line_count=line_count, char_count=char_count,
                    )

        return CodeValidationResult(valid=True, line_count=line_count, char_count=char_count)

    def _execute_code(self, doc: Any, code: str, timeout: int = 30) -> CodeExecutionResult:
        """
        Execute validated python-docx code in a restricted sandbox.

        The code runs in exec() with a namespace containing only `doc`,
        docx types, and safe builtins. A daemon thread enforces the timeout.
        """
        # --- Count before ---
        try:
            paragraphs_before = len(doc.paragraphs)
            runs_before = sum(len(p.runs) for p in doc.paragraphs)
        except Exception:
            paragraphs_before = runs_before = 0

        # --- Build restricted namespace ---
        namespace: dict = {"doc": doc}

        try:
            from docx.shared import RGBColor, Pt, Inches, Cm, Emu
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
            from docx.enum.style import WD_STYLE_TYPE
            from docx.enum.dml import WD_COLOR_INDEX
            namespace.update({
                "RGBColor": RGBColor, "Pt": Pt, "Inches": Inches,
                "Cm": Cm, "Emu": Emu,
                "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
                "WD_LINE_SPACING": WD_LINE_SPACING,
                "WD_BREAK": WD_BREAK,
                "WD_STYLE_TYPE": WD_STYLE_TYPE,
                "WD_COLOR_INDEX": WD_COLOR_INDEX,
            })
        except ImportError as e:
            logger.warning(f"Could not import some docx types into sandbox: {e}")

        namespace["__builtins__"] = {
            "range": range, "len": len, "str": str, "int": int,
            "float": float, "bool": bool, "list": list, "dict": dict,
            "set": set, "tuple": tuple, "enumerate": enumerate,
            "zip": zip, "min": min, "max": max, "round": round,
            "abs": abs, "sorted": sorted, "reversed": reversed,
            "isinstance": isinstance, "hasattr": hasattr,
            "True": True, "False": False, "None": None,
            "print": lambda *a, **kw: None,
        }

        # --- Thread-based execution with timeout ---
        exec_error: list = []

        def _run():
            try:
                exec(compile(code, "<ai_generated>", "exec"), namespace)
            except Exception as exc:
                exec_error.append(exc)

        thread = threading.Thread(target=_run, daemon=True)
        thread.start()
        thread.join(timeout=timeout)

        if thread.is_alive():
            return CodeExecutionResult(
                success=False,
                error=f"Execution timed out after {timeout} seconds",
                paragraphs_before=paragraphs_before, runs_before=runs_before,
            )

        if exec_error:
            exc = exec_error[0]
            return CodeExecutionResult(
                success=False,
                error=f"{type(exc).__name__}: {exc}",
                paragraphs_before=paragraphs_before, runs_before=runs_before,
            )

        # --- Count after ---
        try:
            paragraphs_after = len(doc.paragraphs)
            runs_after = sum(len(p.runs) for p in doc.paragraphs)
        except Exception:
            paragraphs_after = paragraphs_before
            runs_after = runs_before

        run_delta = abs(runs_after - runs_before)
        para_delta = abs(paragraphs_after - paragraphs_before)

        return CodeExecutionResult(
            success=True,
            message=f"Code executed: {run_delta} run(s) and {para_delta} paragraph(s) affected",
            paragraphs_before=paragraphs_before, runs_before=runs_before,
            paragraphs_after=paragraphs_after, runs_after=runs_after,
        )

    # ------------------------------------------------------------------
    # Session helpers
    # ------------------------------------------------------------------

    @staticmethod
    def prepare_session(original_path: str) -> Optional[str]:
        """
        Copy the original file to a temp location for session-mode editing.

        Args:
            original_path: Path to the original .docx file.

        Returns:
            Path to the temp copy, or None on failure.
        """
        try:
            src = Path(original_path)
            if not src.exists():
                logger.error(f"Cannot prepare session: file not found: {src}")
                return None

            # Use %APPDATA%/Specter/temp/ so files are easy to find
            appdata = os.environ.get("APPDATA", "")
            if appdata:
                temp_dir = Path(appdata) / "Specter" / "temp"
            else:
                temp_dir = Path(tempfile.gettempdir()) / "specter_sessions"
            temp_dir.mkdir(parents=True, exist_ok=True)

            temp_path = temp_dir / src.name
            # If a previous session temp exists, remove it
            if temp_path.exists():
                temp_path.unlink()

            shutil.copy2(str(src), str(temp_path))
            logger.info(f"Session temp copy: {src} → {temp_path}")
            return str(temp_path)

        except Exception as e:
            logger.error(f"Failed to prepare session temp copy: {e}", exc_info=True)
            return None

    @staticmethod
    def finalize_session(temp_path: str, original_path: str, auto_open: bool = True) -> Optional[str]:
        """
        Save the session result to AppData temp dir and clean up.

        Creates ``<name>_formatted.docx`` in AppData/Specter/temp/ and
        optionally opens it with the default application.

        Args:
            temp_path: Path to the temp working copy.
            original_path: Path to the original file.
            auto_open: Whether to open the formatted file after saving.

        Returns:
            Path to the final saved file, or None on failure.
        """
        try:
            src = Path(temp_path)
            orig = Path(original_path)

            if not src.exists():
                logger.error(f"Temp file not found: {src}")
                return None

            # Save to AppData temp dir (not next to original)
            final_path = DocxFormatterSkill._get_output_path(orig)

            # Check file lock and find alternative if needed
            if DocxFormatterSkill._is_file_locked(final_path):
                final_path = DocxFormatterSkill._find_available_path(final_path)

            try:
                shutil.copy2(str(src), str(final_path))
            except PermissionError:
                final_path = DocxFormatterSkill._find_available_path(final_path)
                shutil.copy2(str(src), str(final_path))

            logger.info(f"Session finalized: {final_path}")

            # Clean up temp working copy
            try:
                src.unlink()
            except Exception:
                pass

            # Open the formatted file
            if auto_open:
                DocxFormatterSkill.open_file(str(final_path))

            return str(final_path)

        except Exception as e:
            logger.error(f"Failed to finalize session: {e}", exc_info=True)
            return None

    @staticmethod
    def cancel_session(temp_path: str) -> None:
        """Discard the temp copy without saving."""
        try:
            p = Path(temp_path)
            if p.exists():
                p.unlink()
                logger.info(f"Session cancelled, temp removed: {p}")
        except Exception as e:
            logger.warning(f"Failed to remove temp file: {e}")

    # ------------------------------------------------------------------
    # Lifecycle hooks
    # ------------------------------------------------------------------

    async def on_success(self, result: SkillResult) -> None:
        """Log successful document formatting."""
        logger.info(f"Document formatter succeeded: {result.message}")

    async def on_error(self, result: SkillResult) -> None:
        """Log document formatting failure."""
        logger.warning(f"Document formatter failed: {result.error}")
