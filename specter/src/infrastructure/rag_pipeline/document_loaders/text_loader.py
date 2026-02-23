"""
Text Document Loader

Handles various text-based document formats:
- Plain text files (.txt, .md, .rst)
- Word documents (.docx)
- Rich text format (.rtf)
- HTML files (.html, .htm)
- Source code files
- Encoding detection and handling
"""

import asyncio
import logging
import time
from pathlib import Path
from typing import Dict, List, Optional, Union, Any

try:
    import chardet
    CHARDET_AVAILABLE = True
except ImportError:
    CHARDET_AVAILABLE = False
    chardet = None

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    docx = None

try:
    from bs4 import BeautifulSoup
    import html2text
    HTML_PROCESSING_AVAILABLE = True
except ImportError:
    HTML_PROCESSING_AVAILABLE = False
    BeautifulSoup = None
    html2text = None

from .base_loader import BaseDocumentLoader, Document, DocumentMetadata, DocumentLoadError

logger = logging.getLogger("specter.text_loader")


class TextLoader(BaseDocumentLoader):
    """
    Comprehensive text document loader supporting multiple formats.
    
    Features:
    - Multiple text formats support
    - Automatic encoding detection
    - Metadata extraction from content
    - HTML to text conversion
    - Word document processing
    - Source code handling
    """
    
    SUPPORTED_EXTENSIONS = {
        # Plain text
        '.txt', '.text', '.md', '.markdown', '.rst', '.log',
        # Documents
        '.docx', '.rtf', '.odt',
        # Web formats
        '.html', '.htm', '.xml',
        # Source code (common types)
        '.py', '.js', '.java', '.cpp', '.c', '.h', '.cs', '.go', '.rs',
        '.php', '.rb', '.scala', '.swift', '.kt', '.ts', '.jsx', '.tsx',
        '.json', '.yaml', '.yml', '.toml', '.ini', '.cfg', '.conf',
        '.sql', '.sh', '.bash', '.zsh', '.ps1', '.bat', '.cmd'
    }
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize text loader.
        
        Args:
            config: Configuration options including:
                - encoding_detection: Enable automatic encoding detection (default: True)
                - fallback_encoding: Encoding to use if detection fails (default: 'utf-8')
                - html_to_text: Convert HTML to clean text (default: True)
                - preserve_code_structure: Preserve indentation in code files (default: True)
                - extract_metadata_from_content: Extract title, author etc. from content (default: True)
        """
        super().__init__(config)
        
        self.encoding_detection = self.config.get('encoding_detection', True)
        self.fallback_encoding = self.config.get('fallback_encoding', 'utf-8')
        self.html_to_text = self.config.get('html_to_text', True)
        self.preserve_code_structure = self.config.get('preserve_code_structure', True)
        self.extract_metadata_from_content = self.config.get('extract_metadata_from_content', True)
        
        # Warn about missing optional dependencies
        if not CHARDET_AVAILABLE and self.encoding_detection:
            logger.warning("chardet not available - encoding detection disabled")
            self.encoding_detection = False
        
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available - Word document support disabled")
        
        if not HTML_PROCESSING_AVAILABLE and self.html_to_text:
            logger.warning("BeautifulSoup/html2text not available - HTML processing disabled")
            self.html_to_text = False
    
    def supports(self, source: Union[str, Path]) -> bool:
        """Check if this loader supports the given source."""
        if isinstance(source, str) and source.startswith(('http://', 'https://')):
            return False  # URLs handled by web loader
        
        path = Path(source)
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    async def load(self, source: Union[str, Path]) -> Document:
        """
        Load text document.
        
        Args:
            source: Path to text file
            
        Returns:
            Document with extracted text and metadata
        """
        start_time = time.time()
        path = Path(source)
        
        # Validate source
        self.validate_source(path)
        
        # Get basic metadata
        metadata = self.get_metadata_from_path(path)
        metadata.loader_type = self.__class__.__name__
        
        try:
            # Choose loading method based on file extension
            extension = path.suffix.lower()
            
            if extension == '.docx' and DOCX_AVAILABLE:
                content = await self._load_docx(path, metadata)
            elif extension in ['.html', '.htm'] and self.html_to_text:
                content = await self._load_html(path, metadata)
            elif extension == '.rtf':
                content = await self._load_rtf(path, metadata)
            else:
                content = await self._load_text(path, metadata)
            
            # Extract additional metadata from content
            if self.extract_metadata_from_content:
                self._extract_content_metadata(content, metadata)
            
            # Clean content
            if not self._is_code_file(path):
                content = self.clean_content(content)
            elif self.preserve_code_structure:
                # For code files, only do minimal cleaning
                content = content.strip()
            
            if not content or len(content.strip()) < self.config.get('min_content_length', 10):
                raise DocumentLoadError("No meaningful content extracted", str(path))
            
            # Extract title if not already set
            if not metadata.title:
                metadata.title = self.extract_title(content, metadata)
            
            # Detect language for non-code files
            if not self._is_code_file(path):
                metadata.language = self.detect_language(content)
            else:
                metadata.language = self._detect_code_language(path)
            
            # Record processing time
            processing_time = time.time() - start_time
            metadata.processing_time = processing_time
            
            # Update statistics
            self._stats['documents_loaded'] += 1
            self._stats['total_processing_time'] += processing_time
            self._stats['total_content_size'] += len(content)
            
            document = Document(content=content, metadata=metadata)
            self.logger.info(f"Successfully loaded text: {path} ({len(content)} chars, {processing_time:.2f}s)")
            
            return document
            
        except Exception as e:
            self._stats['documents_failed'] += 1
            if isinstance(e, DocumentLoadError):
                raise
            else:
                raise DocumentLoadError(f"Text extraction failed: {str(e)}", str(path), e)
    
    async def _load_text(self, path: Path, metadata: DocumentMetadata) -> str:
        """Load plain text file with encoding detection."""
        def load_sync():
            # Try to detect encoding
            encoding = self.fallback_encoding
            
            if self.encoding_detection and CHARDET_AVAILABLE:
                try:
                    with open(path, 'rb') as f:
                        raw_data = f.read()
                        if raw_data:
                            detected = chardet.detect(raw_data)
                            if detected['encoding'] and detected['confidence'] > 0.7:
                                encoding = detected['encoding']
                                logger.debug(f"Detected encoding {encoding} with confidence {detected['confidence']}")
                except Exception as e:
                    logger.warning(f"Encoding detection failed: {e}")
            
            # Load file with detected/fallback encoding
            try:
                with open(path, 'r', encoding=encoding) as f:
                    content = f.read()
                metadata.encoding = encoding
                return content
            except UnicodeDecodeError:
                # Try common encodings as fallback
                fallback_encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                for fallback in fallback_encodings:
                    if fallback != encoding:
                        try:
                            with open(path, 'r', encoding=fallback) as f:
                                content = f.read()
                            metadata.encoding = fallback
                            logger.warning(f"Used fallback encoding {fallback} for {path}")
                            return content
                        except UnicodeDecodeError:
                            continue
                
                # If all else fails, read as binary and decode with errors='replace'
                with open(path, 'rb') as f:
                    raw_content = f.read()
                content = raw_content.decode('utf-8', errors='replace')
                metadata.encoding = 'utf-8 (with errors replaced)'
                logger.warning(f"Used error replacement for {path}")
                return content
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, load_sync)
    
    async def _load_docx(self, path: Path, metadata: DocumentMetadata) -> str:
        """Load Word document (.docx)."""
        def load_sync():
            doc = docx.Document(str(path))
            
            # Extract metadata from document properties
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                if props.title:
                    metadata.title = props.title
                if props.author:
                    metadata.author = props.author
                if props.subject:
                    metadata.subject = props.subject
                if props.created:
                    metadata.created_at = props.created
                if props.modified:
                    metadata.modified_at = props.modified
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(' | '.join(row_text))
            
            content = '\n\n'.join(paragraphs)
            metadata.extraction_method = 'python-docx'
            
            return content
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, load_sync)
    
    async def _load_html(self, path: Path, metadata: DocumentMetadata) -> str:
        """Load HTML file and convert to text."""
        def load_sync():
            # First load the HTML content
            with open(path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            if not HTML_PROCESSING_AVAILABLE:
                # Fallback: just remove tags using simple regex
                import re
                text = re.sub(r'<[^>]+>', '', html_content)
                text = re.sub(r'\s+', ' ', text)
                return text.strip()
            
            # Parse HTML with BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract metadata from HTML head
            if soup.title:
                metadata.title = soup.title.get_text().strip()
            
            # Look for meta tags
            for meta in soup.find_all('meta'):
                name = meta.get('name', '').lower()
                content = meta.get('content', '')
                
                if name in ['author', 'creator'] and content:
                    metadata.author = content
                elif name in ['description', 'subject'] and content:
                    metadata.subject = content
                elif name == 'language' and content:
                    metadata.language = content
            
            # Convert to clean text using html2text
            h = html2text.HTML2Text()
            h.ignore_links = False
            h.ignore_images = True
            h.ignore_emphasis = False
            h.body_width = 0  # Don't wrap lines
            
            text = h.handle(str(soup))
            metadata.extraction_method = 'html2text'
            
            return text
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, load_sync)
    
    async def _load_rtf(self, path: Path, metadata: DocumentMetadata) -> str:
        """Load RTF file (basic implementation)."""
        def load_sync():
            # This is a very basic RTF parser - for production use, consider striprtf library
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Remove RTF control words and groups
            import re
            
            # Remove RTF control words
            content = re.sub(r'\\[a-z]+\d*\s?', '', content)
            # Remove RTF groups
            content = re.sub(r'[{}]', '', content)
            # Clean up multiple spaces
            content = re.sub(r'\s+', ' ', content)
            
            metadata.extraction_method = 'basic_rtf_parser'
            return content.strip()
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, load_sync)
    
    def _extract_content_metadata(self, content: str, metadata: DocumentMetadata) -> None:
        """Extract metadata from content (title, sections, etc.)."""
        lines = content.split('\n')
        
        # Look for markdown-style headers
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if line.startswith('#'):
                # Markdown header
                title = line.lstrip('# ').strip()
                if title and not metadata.title:
                    metadata.title = title
                break
            elif line.startswith('=') or line.startswith('-'):
                # Could be RST-style header
                continue
        
        # Extract sections/structure information
        sections = []
        for line in lines:
            line = line.strip()
            if line.startswith('#'):
                sections.append(line.lstrip('# ').strip())
        
        if sections:
            if not metadata.custom:
                metadata.custom = {}
            metadata.custom['sections'] = sections
    
    def _is_code_file(self, path: Path) -> bool:
        """Check if the file is a source code file."""
        code_extensions = {
            '.py', '.js', '.java', '.cpp', '.c', '.h', '.cs', '.go', '.rs',
            '.php', '.rb', '.scala', '.swift', '.kt', '.ts', '.jsx', '.tsx',
            '.json', '.yaml', '.yml', '.toml', '.sql', '.sh', '.bash', '.zsh'
        }
        return path.suffix.lower() in code_extensions
    
    def _detect_code_language(self, path: Path) -> Optional[str]:
        """Detect programming language from file extension."""
        extension_to_language = {
            '.py': 'python',
            '.js': 'javascript',
            '.jsx': 'javascript',
            '.ts': 'typescript',
            '.tsx': 'typescript',
            '.java': 'java',
            '.cpp': 'cpp',
            '.c': 'c',
            '.h': 'c',
            '.cs': 'csharp',
            '.go': 'go',
            '.rs': 'rust',
            '.php': 'php',
            '.rb': 'ruby',
            '.scala': 'scala',
            '.swift': 'swift',
            '.kt': 'kotlin',
            '.json': 'json',
            '.yaml': 'yaml',
            '.yml': 'yaml',
            '.toml': 'toml',
            '.sql': 'sql',
            '.sh': 'bash',
            '.bash': 'bash',
            '.zsh': 'zsh',
            '.ps1': 'powershell',
            '.bat': 'batch',
            '.cmd': 'batch'
        }
        
        return extension_to_language.get(path.suffix.lower())
    
    def clean_content(self, content: str) -> str:
        """Override clean_content for text-specific cleaning."""
        if not content:
            return ""
        
        # Standard cleaning
        content = super().clean_content(content)
        
        # Text-specific cleaning
        if self.config.get('remove_extra_newlines', True):
            # Replace multiple consecutive newlines with at most two
            import re
            content = re.sub(r'\n{3,}', '\n\n', content)
        
        # Remove control characters (except newlines and tabs)
        if self.config.get('remove_control_chars', True):
            import re
            content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', content)
        
        return content